"""Parser for DOCX files using python-docx."""

from pathlib import Path


def parse_docx(path: Path) -> str:
    """Extract text from a DOCX file preserving heading hierarchy and lists.

    Converts DOCX structure to markdown-style formatting.
    """
    from docx import Document

    doc = Document(str(path))
    sections: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            sections.append("")
            continue

        style_name = (para.style.name or "").lower()

        # Map heading styles to markdown
        if "heading 1" in style_name:
            sections.append(f"# {text}")
        elif "heading 2" in style_name:
            sections.append(f"## {text}")
        elif "heading 3" in style_name:
            sections.append(f"### {text}")
        elif "heading 4" in style_name:
            sections.append(f"#### {text}")
        elif "title" in style_name:
            sections.append(f"# {text}")
        elif "subtitle" in style_name:
            sections.append(f"## {text}")
        elif "list" in style_name or "bullet" in style_name:
            sections.append(f"- {text}")
        elif "quote" in style_name:
            sections.append(f"> {text}")
        else:
            sections.append(text)

    # Extract tables
    for table in doc.tables:
        table_lines: list[str] = []
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            table_lines.append("| " + " | ".join(cells) + " |")
            if row_idx == 0:
                table_lines.append("|" + "|".join(["---"] * len(cells)) + "|")
        if table_lines:
            sections.append("\n" + "\n".join(table_lines) + "\n")

    result = "\n".join(sections).strip()
    if not result:
        raise ValueError(f"No text content could be extracted from {path}")
    return result
